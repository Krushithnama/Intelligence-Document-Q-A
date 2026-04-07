from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    sections: list[tuple[str, str]]  # (section_title, section_text)


def _normalize_ws(s: str) -> str:
    return "\n".join([line.rstrip() for line in s.replace("\r\n", "\n").split("\n")]).strip()


def parse_txt(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore")
    text = _normalize_ws(text)
    return ParsedDocument(text=text, sections=[("Document", text)])


def parse_pdf(path: Path) -> ParsedDocument:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        extracted = page.extract_text() or ""
        extracted = _normalize_ws(extracted)
        pages.append((f"Page {i+1}", extracted))
    full = "\n\n".join([t for _, t in pages]).strip()
    return ParsedDocument(text=full, sections=pages if pages else [("Document", full)])


def parse_docx(path: Path) -> ParsedDocument:
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = _normalize_ws("\n".join(paragraphs))
    return ParsedDocument(text=text, sections=[("Document", text)])


def parse_file(path: Path, content_type: str | None = None) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix in (".docx",):
        return parse_docx(path)
    if suffix in (".txt",):
        return parse_txt(path)
    # fallback by content type
    if content_type:
        if "pdf" in content_type:
            return parse_pdf(path)
        if "word" in content_type or "docx" in content_type:
            return parse_docx(path)
        if "text" in content_type:
            return parse_txt(path)
    raise ValueError(f"Unsupported file type: {path.name}")

